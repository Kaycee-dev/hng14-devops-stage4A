"""
Generate docs/interview_prep.docx from the markdown source.
Run from repo root: python scripts/build_interview_prep_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
MD_FILE = ROOT / "docs" / "interview_prep.md"
OUT_FILE = ROOT / "docs" / "interview_prep.docx"


# ── colour palette (stored as hex strings for XML; RGBColor built on the fly) ──
HEADING1_HEX   = "1A1A2E"
HEADING2_HEX   = "16213E"
HEADING3_HEX   = "0F3A5C"
ACCENT_HEX     = "1B6CA8"
TABLE_HDR_HEX  = "1A1A2E"
TABLE_ALT_HEX  = "EFF3F8"

def _rgb(hex6: str) -> RGBColor:
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return RGBColor(r, g, b)

HEADING1_COLOR  = _rgb(HEADING1_HEX)
HEADING2_COLOR  = _rgb(HEADING2_HEX)
HEADING3_COLOR  = _rgb(HEADING3_HEX)
ACCENT_COLOR    = _rgb(ACCENT_HEX)
TABLE_HEADER_BG = _rgb(TABLE_HDR_HEX)
TABLE_ROW_ALT   = _rgb(TABLE_ALT_HEX)


def set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_para_border_left(para, hex_color: str) -> None:
    """Thin left border — used for blockquote / tip paragraphs."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)
    para.paragraph_format.left_indent = Cm(0.6)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for level, (name, size, color, bold) in enumerate([
        ("Heading 1", 20, HEADING1_COLOR, True),
        ("Heading 2", 15, HEADING2_COLOR, True),
        ("Heading 3", 12, HEADING3_COLOR, True),
    ], start=1):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def add_code_block(doc: Document, code: str) -> None:
    """Add a shaded monospace code block."""
    style = doc.styles["Normal"]
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        # shade the paragraph cell via pPr shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)
    # small gap after code block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(4)


def apply_inline_formatting(run_parent, text: str) -> None:
    """
    Parse inline markdown in `text` and add formatted runs to run_parent.
    Handles **bold**, `code`, and plain text segments.
    """
    # Combined pattern: backtick code OR bold (**...**)
    TOKEN = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*)')
    parts = TOKEN.split(text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = run_parent.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = ACCENT_COLOR
        elif part.startswith("**") and part.endswith("**"):
            run = run_parent.add_run(part[2:-2])
            run.font.bold = True
        else:
            if part:
                run_parent.add_run(part)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    apply_inline_formatting(p, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            if i == 0:
                set_cell_bg(cell, TABLE_HDR_HEX)
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            else:
                if i % 2 == 0:
                    set_cell_bg(cell, TABLE_ALT_HEX)
                p = cell.paragraphs[0]
                apply_inline_formatting(p, cell_text)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B6CA8")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_and_render(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ─────────────────────────────────────────────────
        if line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────────
        if line.startswith("|") and "|" in line:
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Filter out separator rows (|---|---|)
            rows: list[list[str]] = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.fullmatch(r"[-: ]+", c) for c in cells if c):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[*\-] (.+)$", line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=indent)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Q&A pair detection ("**Q: ..." / "A: ...") ───────────────────────
        if line.startswith("**Q:"):
            # Render question in a distinct style
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            set_para_border_left(p, ACCENT_HEX)
            run = p.add_run(line.strip("*").strip())
            run.font.bold = True
            run.font.color.rgb = ACCENT_COLOR
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            i += 1
            continue

        if line.startswith("A: "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.6)
            apply_inline_formatting(p, line[3:])
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        apply_inline_formatting(p, line)
        i += 1


def build_cover(doc: Document) -> None:
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SwiftDeploy Stage 4B")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = HEADING1_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Interview Prep & Proof-of-Work Pack")
    run2.font.name = "Calibri"
    run2.font.size = Pt(16)
    run2.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run("HNG DevOps Track — Stage 4B Submission\nKelechi Uba  ·  2026")
    run3.font.name = "Calibri"
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_horizontal_rule(doc)
    doc.add_page_break()


def main() -> None:
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    configure_styles(doc)
    build_cover(doc)

    md_text = MD_FILE.read_text(encoding="utf-8")
    # Strip the H1 title from the markdown (already on cover page)
    md_text = re.sub(r"^# .+\n", "", md_text, count=1)
    parse_and_render(doc, md_text)

    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    main()
